from docx import Document #Dùng để làm việc với các file docx
'''
ví dụ
from docx import Document
# Mở file Word
doc = Document("example.docx")
# Duyệt qua từng đoạn văn trong tài liệu
for paragraph in doc.paragraphs:
    print(paragraph.text)
'''
import pandas as pd
'''
Dùng để thao tác với dữ liệu dạng bảng (DataFrame)
Ví dụ:
import pandas as pd
# Dữ liệu mẫu
data = {
    "Tên sản phẩm": ["Laptop", "Điện thoại", "Máy tính bảng"],
    "Giá": [15000000, 7000000, 10000000],
    "Số lượng": [10, 20, 15]
}
# Tạo DataFrame từ dictionary
df = pd.DataFrame(data)
# Hiển thị DataFrame
print(df)
'''

def read_data_from_word(file_path):
    doc = Document(file_path)
    '''
    Đối tượng Document chứa tất cả thông tin của tài liệu, bao gồm
    các đoạn văn, các bảng, các tiêu đề, hình ảnh,...
    '''
    data = []
    current_disease = {} #Tạo một dict trống

    for paragraph in doc.paragraphs:
        '''
        Lọc từng đoạn văn trong doc.paragraphs,
        text = paragraph.text.strip() nghĩa là paragraph.text là nội dung đoạn văn
        .strip() bỏ nghĩa là bỏ khoảng trống(space hoặc kí tự xuống dòng) 
        ở đầu và cuối đoạn văn.
        '''
        text = paragraph.text.strip()

        '''
        Câu lệnh:
        current_disease = {"Disease": text.replace("Tên bệnh:", "").strip()}
        Giải thích chi tiết
        Câu lệnh này:
        Khởi tạo một dictionary mới tên là current_disease.
        Thêm cặp key-value đầu tiên vào dictionary:
        Key: "Disease"
        Value: Chuỗi văn bản (tên bệnh) được lấy từ biến text, sau khi xử lý loại bỏ phần tiền tố "Tên bệnh:" và các khoảng trắng thừa.
        Phân tích từng phần
        1. Tạo dictionary mới
        current_disease = {...}
        current_disease: Biến này lưu trữ thông tin về một bệnh cụ thể.
        {}: Sử dụng cú pháp dictionary để khởi tạo một dictionary mới.
        2. Key của dictionary
        "Disease"
        Đây là khóa (key) của dictionary, được sử dụng để xác định giá trị (value) tương ứng với tên bệnh.
        3. Giá trị của key
        text.replace("Tên bệnh:", "").strip()
        3.1. text
        text là nội dung văn bản của đoạn văn hiện tại trong file Word, được lấy từ paragraph.text.
        Ví dụ: Nếu đoạn văn là "Tên bệnh: Cúm", thì text = "Tên bệnh: Cúm".
        3.2. .replace("Tên bệnh:", "")
        .replace("Tên bệnh:", ""): Xóa đi phần tiền tố "Tên bệnh:" trong chuỗi.
        Ví dụ:
        text = "Tên bệnh: Cúm"
        text.replace("Tên bệnh:", "")
        Kết quả:
        " Cúm"
        3.3. .strip()
        .strip(): Loại bỏ các khoảng trắng thừa ở đầu và cuối chuỗi.
        Ví dụ:
        " Cúm".strip()
        Kết quả:
        "Cúm"
        '''
        if text.startswith("Tên bệnh:"):
            if current_disease:  # Lưu bệnh trước đó
                data.append(current_disease)
            current_disease = {"Disease": text.replace("Tên bệnh:", "").strip()}
        elif text.startswith("Triệu chứng:"):
            current_disease["Symptoms"] = text.replace("Triệu chứng:", "").strip()
        elif text.startswith("Nhóm nguy cơ:"):
            current_disease["Risk Group"] = text.replace("Nhóm nguy cơ:", "").strip()
        elif text.startswith("Giai đoạn:"):
            current_disease["Stage"] = text.replace("Giai đoạn:", "").strip()
        elif text.startswith("Phòng ngừa:"):
            current_disease["Prevention"] = text.replace("Phòng ngừa:", "").strip()

    if current_disease:  # Thêm bệnh cuối cùng
        data.append(current_disease)
    print(data)

    return pd.DataFrame(data)


# Đọc dữ liệu từ file Word
file_path = "disease_data.docx"
df = read_data_from_word(file_path)

# Lưu dữ liệu thành file CSV
df.to_csv("data/disease_data.csv", index=False)
print("Dữ liệu đã được lưu vào data/disease_data.csv")



