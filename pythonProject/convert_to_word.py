from docx import Document
import pandas as pd

# Đọc dữ liệu từ CSV
df = pd.read_csv("disease_data_crawled.csv")

# Tạo file Word
doc = Document()

for _, row in df.iterrows():
    doc.add_heading(f"Tên bệnh: {row['Tên bệnh']}", level=1)
    doc.add_paragraph(f"Nguyên nhân: {row['Nguyên nhân']}")
    doc.add_paragraph(f"Triệu chứng: {row['Triệu chứng']}")
    doc.add_paragraph(f"Điều trị: {row['Điều trị']}")
    doc.add_paragraph("\n")

# Lưu file Word
doc.save("disease_data.docx")
print("Dữ liệu đã được lưu vào 'disease_data.docx'")
